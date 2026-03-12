
import os
import streamlit as st
from groq import Groq
import requests
import folium
from streamlit_folium import st_folium
from geopy.geocoders import Nominatim
from docx import Document
import io


#  Live Weather Fetcher 
def get_live_weather(city_name):
    try:
        api_key = st.secrets["WEATHER_API_KEY"]
    except KeyError:
        return "Weather API key not found."

    base_url = "http://api.openweathermap.org/data/2.5/weather"
    params = {
        "q": city_name,
        "appid": api_key,
        "units": "metric" 
    }
    
    try:
        response = requests.get(base_url, params=params)
        data = response.json()
        
        if data["cod"] == 200:
            temp = round(data["main"]["temp"])
            description = data["weather"][0]["description"].title()
            return f"**Live Weather in {city_name.title()}:** {temp}°C, {description} 🌤️"
        else:
            return None 
    except Exception as e:
        return None

# --- Groq API Setup ---
def get_groq_client():
    api_key = None
    try:
        api_key = st.secrets["GROQ_API_KEY"]
    except (KeyError, FileNotFoundError):
        api_key = os.environ.get('GROQ_API_KEY')
        
    if not api_key:
        st.error("GROQ_API_KEY not found. Please set it in your Streamlit secrets.")
        return None
    return Groq(api_key=api_key)

# MAP
def get_coordinates(city_name):
    """Converts a city name into latitude and longitude."""
    try:
        # Using Nominatim geocoder to get coordinates for the city
        geolocator = Nominatim(user_agent="ai_travel_planner_app")
        location = geolocator.geocode(city_name)
        if location:
            return location.latitude, location.longitude
        return None
    except Exception as e:
        return None

def display_city_map(lat, lon, city_name):
    """Generates and displays an interactive Folium map."""
    # Create a map centered on the coordinates
    city_map = folium.Map(location=[lat, lon], zoom_start=11)
    
    # Add a marker for the destination
    folium.Marker(
        [lat, lon], 
        popup=city_name,
        tooltip=f"Explore {city_name}!"
    ).add_to(city_map)
    
    # Display it in Streamlit
    st_folium(city_map, width=700, height=400)

# --- Word Document Generation ---
def generate_word_doc(markdown_text, location_name):
    doc = Document()
    doc.add_heading(f'Travel Itinerary: {location_name.title()}', 0)
    
    # A simple parser to clean up Markdown and format it for Word
    for line in markdown_text.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        # Handle Headers
        if line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=2)
        elif line.startswith('**') and line.endswith('**'):
            doc.add_heading(line.replace('**', ''), level=3)
        # Handle Bullet Points
        elif line.startswith('* '):
            clean_line = line.replace('**', '').replace('* ', '')
            doc.add_paragraph(clean_line, style='List Bullet')
        # Handle regular text
        else:
            clean_line = line.replace('**', '')
            doc.add_paragraph(clean_line)
            
    # Save the document to an in-memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# Travel Recommendation 
def generate_travel_recommendation(client, location=None, destination_type=None, interests=None, budget_range=None, travel_companions=None, duration=None, time_of_year=None):
    prompt = f"""
    Generate a personalized travel recommendation for {location or 'any location'} based on these preferences:
    - Destination Type: {destination_type or 'Any'}
    - Interests: {interests or 'General'}
    - Budget: {budget_range or 'Any'}
    - Traveling With: {travel_companions or 'Any'}
    - Trip Duration: {duration or 'Not specified'}
    - Time of Year: {time_of_year or 'Not specified'}
    
    Please provide a detailed and engaging travel plan. Suggest a specific city or region.
    Include recommendations for:
    1. **Accommodation:** Suggest a type of place to stay that fits the budget.
    2. **Activities:** A brief, day-by-day itinerary or a list of 3-5 key activities.
    3. **Food:** Mention 1-2 local dishes or types of restaurants to try.
    4. **Why this destination?** Briefly explain why this location is a great fit.
    Format the output using Markdown for readability. Use headings, bold text, bullet points, and sprinkle in some fun emojis.
    """
    try:
        chat_completion = client.chat.completions.create(
            messages=[
                {
                    "role": "user",
                    "content": prompt,
                }
            ],
            model="llama-3.1-8b-instant", 
        )
        return chat_completion.choices[0].message.content
    except Exception as e:
        st.error(f"An error occurred while generating the recommendation: {e}")
        return "Sorry, I couldn't generate a recommendation at this time. Please try again."

def display_recommendations(recommendation_text):
    st.subheader("Your Personalized Travel Recommendation! ✈️🌟")
    st.markdown(recommendation_text)

# streamlit layoput
st.set_page_config(layout="wide")
st.title("AI Travel Recommendation System 🤖✈️🌴")
st.markdown("Get personalized travel recommendations. Fill out the form below to get started! 🚀")

# Initialize Client
groq_client = get_groq_client()

if groq_client:
    st.header("Your Travel Preferences 🧳")

    
    popular_locations = ["Paris", "Tokyo", "Sydney", "New York", "London", "Other"]
    location_choice = st.selectbox("🌍 Choose your destination:", popular_locations)
    
    if location_choice == "Other":
        location = st.text_input(
            "Enter your destination (City or Region)",
            placeholder="e.g., Chennai, Cape Town, Rio de Janeiro",
            key="custom_location_input"
        )
    else:
        location = location_choice

    
    if location:
        weather_info = get_live_weather(location)
        if weather_info:
            st.info(weather_info)
            
        
        coords = get_coordinates(location)
        if coords:
            lat, lon = coords
            st.markdown(f"### 🗺️ Map of {location.title()}")
            display_city_map(lat, lon, location.title())
    # -------------------------------------

    col1, col2 = st.columns(2) 
    with col1:
        destination_type = st.selectbox(
            "🏞️ What type of destination are you interested in?",
            ["Any", "City", "Beach", "Mountains", "Rural", "Adventure", "Historical", "Relaxation"]
        )
        interests = st.multiselect(
            "🎯 What are your interests?",
            ["History", "Adventure", "Relaxation", "Food", "Art", "Nature", "Nightlife", "Shopping"]
        )
        budget_range = st.select_slider(
            "💸 What is your budget range?",
            options=["Budget", "Mid-range", "Luxury"]
        )
    with col2:
        travel_companions = st.selectbox(
            "👫 Who are you traveling with?",
            ["Any", "Solo", "Family", "Couple", "Friends"]
        )
        duration = st.text_input(
            "⏳ How long do you plan to travel? (e.g., 7 days, 1 week)"
        )
        time_of_year = st.text_input(
            "📅 When do you plan to travel? (e.g., June, Summer, Winter)"
        )


    # Button to trigger recommendation generation
    if st.button("Get Recommendations 🎉"):
        selected_destination_type = destination_type if destination_type != "Any" else None
        selected_travel_companions = travel_companions if travel_companions != "Any" else None
        formatted_interests = ", ".join(interests) if interests else None
        
        with st.spinner("✨ Generating your personalized itinerary..."):
            recommendations = generate_travel_recommendation(
                client=groq_client,
                location=location or None,
                destination_type=selected_destination_type,
                interests=formatted_interests,
                budget_range=budget_range,
                travel_companions=selected_travel_companions,
                duration=duration or None,
                time_of_year=time_of_year or None
            )
            
        display_recommendations(recommendations)
        
       
        # We replace spaces in the location name with underscores for a clean filename
        safe_location_name = location.replace(" ", "_") if location else "My"
        
        # --- NEW: Word Doc Download Button ---
        safe_location_name = location.replace(" ", "_") if location else "Destination"
        
        # Generate the Word file
        word_file = generate_word_doc(recommendations, location)
        
        st.download_button(
            label="📄 Download as Word Document",
            data=word_file,
            file_name=f"{safe_location_name}_Travel_Plan.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
else:
    st.warning("API is not configured. Please provide your API key to continue.")
